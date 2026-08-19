# -*- coding: utf-8 -*-
"""
通用 md → 排版 docx 生成（教材目录/正文章节）
与 docx-template-fill 的 fill_docx_template.py 区别：本脚本从零生成格式化 docx（标题黑体/正文宋体五号/1.5倍行距/表格），
不是填入现成模板表格。

用法（PowerShell，勿用 & 调用，直接全路径）：
  D:\4.1\python\bin\python.exe md_to_docx.py --md "目录.md" --out "目录.docx"

支持的 md 结构：
  # 文档标题（16pt 黑体居中）
  ## 一级标题（14pt 黑体）
  ### 第X章 xxx / ### 其他（12pt 黑体，章级）
  1.1 节 / 1.1.1 小节 / A.1 附录节（11/10.5pt 黑体，按编号识别）
  | 表格（Table Grid，首行表头黑体加粗居中）
  > 引用 / - 列表 / **加粗** / ``` 代码块 / <!-- 注释

已验证：2026-08-18 目录_2026_08_18_资料融合修订版.md → docx（14章+学时表+更新点表格）
"""
import argparse
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_font(run, name_cn, size, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if level == 0:
        pf.line_spacing = 1.5; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.space_after = Pt(12)
        run = p.add_run(text); set_font(run, "黑体", 16, bold=True)
    elif level == 1:
        pf.line_spacing = 1.5; pf.space_before = Pt(12); pf.space_after = Pt(6)
        run = p.add_run(text); set_font(run, "黑体", 14, bold=True)
    elif level == 2:
        pf.line_spacing = 1.5; pf.space_before = Pt(6); pf.space_after = Pt(3)
        run = p.add_run(text); set_font(run, "黑体", 12, bold=True)
    elif level == 3:
        pf.line_spacing = 1.5
        run = p.add_run(text); set_font(run, "黑体", 11, bold=True)
    else:
        pf.line_spacing = 1.5
        run = p.add_run(text); set_font(run, "黑体", 10.5, bold=True)


def add_body(doc, text, is_bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_font(run, "宋体", 10.5, bold=True)
        elif part:
            run = p.add_run(part); set_font(run, "宋体", 10.5, bold=is_bold)


def add_list(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.74)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    first = True
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_font(run, "宋体", 10.5, bold=True)
        elif part:
            prefix = "• " if first else ""
            run = p.add_run(prefix + part); set_font(run, "宋体", 10.5)
            first = False


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if ri == 0:
                set_font(run, "黑体", 10.5, bold=True); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_font(run, "宋体", 10.5)


def parse_table(lines):
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln or re.match(r"^[\|\s:\-]+$", ln):
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
    return rows


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    with open(args.md, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip("\n").rstrip("\r")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code; i += 1; continue
        if stripped.startswith("<!--"):
            if "-->" not in stripped:
                while i < len(lines) and "-->" not in lines[i]:
                    i += 1
            i += 1; continue
        if stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            add_table(doc, parse_table(tbl_lines)); continue
        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_heading(doc, stripped[2:], 0); i += 1; continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1); i += 1; continue
        m = re.match(r"^###\s+(第\S+章.*)$", stripped)
        if m:
            add_heading(doc, m.group(1), 2); i += 1; continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2); i += 1; continue
        if stripped.startswith("> "):
            add_body(doc, stripped[2:]); i += 1; continue
        if stripped.startswith("- "):
            add_list(doc, stripped[2:]); i += 1; continue
        m = re.match(r"^(\d+\.\d+\.\d+)\s+(.+)$", stripped)
        if m:
            add_heading(doc, f"{m.group(1)}  {m.group(2)}", 4); i += 1; continue
        m = re.match(r"^(\d+\.\d+)\s+(.+)$", stripped)
        if m:
            add_heading(doc, f"{m.group(1)}  {m.group(2)}", 3); i += 1; continue
        m = re.match(r"^(A\.\d+)\s*(.*)$", stripped)
        if m:
            add_heading(doc, f"{m.group(1)}  {m.group(2)}".strip(), 3); i += 1; continue
        if stripped.startswith("**") and stripped.endswith("**"):
            add_body(doc, stripped.strip("*"), is_bold=True); i += 1; continue
        if not stripped:
            i += 1; continue
        add_body(doc, stripped)
        i += 1

    doc.save(args.out)
    print(f"✅ 已生成: {args.out}")


if __name__ == "__main__":
    main()
