# -*- coding: utf-8 -*-
"""
通用 md → docx 模板填充脚本
把按 "## 板块" 分节的 md 内容填入表格式 docx 模板的对应行，另存新文件。

用法（PowerShell，勿用 & 调用，直接全路径）：
  D:\4.1\python\bin\python.exe fill_docx_template.py \
      --template "模板.docx" --content "源.md" --out "输出.docx" \
      --mapping "一、选题思路:8,二、读者对象:9,三、作者情况:10" \
      --top-fields "书名:《Vibe Coding开发与应用>,作者:杨元超"

参数说明：
  --template   模板 docx 绝对路径（必填）
  --content    源 md 绝对路径（必填），按 "## 板块名" 分节
  --out        输出 docx 绝对路径（必填），另存新文件，不覆盖模板
  --mapping    逗号分隔 "板块名:行号"，板块名匹配 md 的 ## 标题
  --top-fields 可选，逗号分隔 "标签:值"，填入顶部单格字段（需在脚本 TOP_ROW 里配好行号+列索引）
  --title-row  可选，书名所在行号（默认 0），填 --title 的值

已验证：2026-08-18《Vibe Coding开发与应用》选题表（人邮社模板 23x15，8 板块 + 书名 + 作者名，20 项验证全过）
"""
import argparse
import re
import sys
from docx import Document

# 模板指导行前缀（给作者的填写说明，正式交付必须剔除；按需扩展）
TEMPLATE_HINTS = (
    "（简要阐述为什么提出此选题方案",
    "（目标读者定位、适合学历",
    "（内容提要，不少于",
    "（根据具体情况列写，资源配套须齐全",
    "（简要分析图书的重点销售区域及销售宣传建议",
    "（以下请选几本销售书",
    "（选几本销售书进行内容",
    "以下文字为示例",
    "参考：本书严格",
    "（以下文字为示例，参考之后请删除",
)

# 顶部单格字段：行号 -> {列索引: 值来源}（--top-fields 传 "标签:值"，标签用于校验）
TOP_ROW_COL = {
    0: {1: "书名"},   # 行0 列1 = 书名
    3: {1: "作者"},   # 行3 列1 = 作者名（注意按 indep 去重后的索引）
}


def read_sections(md_path):
    """按 '## ' 标题切分 md，返回 {板块名: 正文行列表}"""
    with open(md_path, encoding="utf-8") as f:
        content = f.read()
    sections, current, buf = {}, None, []
    for raw in content.splitlines():
        m = re.match(r"^##\s+(.+)$", raw.strip())
        if m:
            if current:
                sections[current] = buf
            current = m.group(1).strip()
            buf = []
        else:
            buf.append(raw)
    if current:
        sections[current] = buf
    return sections


def clean_lines(lines):
    """清洗 markdown → 纯文本段落列表；剔除模板指导行"""
    out = []
    for raw in lines:
        line = raw.strip()
        if not line:
            out.append("")
            continue
        if line == "---":
            continue
        if line.startswith(TEMPLATE_HINTS):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"^>\s*", "", line)
        line = re.sub(r"^[-•]\s+", "", line)
        line = line.strip()
        out.append(line)
    while out and out[0] == "":
        out.pop(0)
    while out and out[-1] == "":
        out.pop()
    return out


def set_cell_text(cell, lines):
    """清空 cell 后写入多段文本（保留 cell 样式）"""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    if lines:
        first.add_run(lines[0])
        for ln in lines[1:]:
            p = cell.add_paragraph()
            p.add_run(ln)


def indep(row):
    """按 _tc 去重返回独立 Cell 对象列表（合并单元格去重）"""
    out = []
    for c in row.cells:
        if not any(c._tc is x for x in [o._tc for o in out]):
            out.append(c)
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--content", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--mapping", required=True, help='板块:行号,板块:行号,...')
    ap.add_argument("--top-fields", default="", help='标签:值,标签:值,...')
    ap.add_argument("--title-row", type=int, default=0)
    ap.add_argument("--title", default="")
    args = ap.parse_args()

    # 板块映射
    mapping = {}
    for item in args.mapping.split(","):
        if ":" in item:
            k, v = item.rsplit(":", 1)
            mapping[k.strip()] = int(v.strip())

    # 顶部字段
    top = {}
    for item in args.top_fields.split(","):
        if ":" in item:
            k, v = item.rsplit(":", 1)
            top[k.strip()] = v.strip()

    sections = read_sections(args.content)
    doc = Document(args.template)
    table = doc.tables[0]

    # 书名
    if args.title:
        row = indep(table.rows[args.title_row])
        set_cell_text(row[1], [args.title])

    # 顶部单格字段（按 TOP_ROW_COL 映射）
    for row_idx, col_map in TOP_ROW_COL.items():
        for col_idx, label in col_map.items():
            if label in top:
                row = indep(table.rows[row_idx])
                set_cell_text(row[col_idx], [top[label]])

    # 板块填充
    filled = []
    for sec_name, row_idx in mapping.items():
        if sec_name not in sections:
            print(f"[SKIP] 未找到板块: {sec_name}")
            continue
        lines = clean_lines(sections[sec_name])
        row = indep(table.rows[row_idx])
        set_cell_text(row[1], lines)
        filled.append((sec_name, len(lines)))

    doc.save(args.out)
    print(f"✅ 已生成: {args.out}")
    for name, n in filled:
        print(f"   {name}: {n} 行")
    print(f"完成 {len(filled)}/{len(mapping)} 板块。")


if __name__ == "__main__":
    main()
