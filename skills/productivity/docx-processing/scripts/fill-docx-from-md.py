# 从 Markdown 填写 Word 模板表格（选题表/申报表/表单类 docx）
# 用法: python fill-docx-from-md.py <模板.docx> <内容.md> <输出.docx> <板块映射.json>
#   <内容.md>   按 "## " 二级标题切分板块，每板块正文填入对应表格行
#   <板块映射.json> 形如 {"一、选题思路": 8, "二、读者对象": 9, ...} (模板表格行号)
# 依赖: python-docx (pip install python-docx)
# 本机 python 不在 PATH: 用 D:\4.1\python\bin\python.exe

import re
import sys
import json
import docx
from docx import Document


def read_sections(md_path):
    """按 '## ' 标题切分 md，返回 {板块名: 正文行列表}"""
    with open(md_path, encoding="utf-8") as f:
        content = f.read()
    sections = {}
    current = None
    buf = []
    for raw in content.splitlines():
        m = re.match(r"^##\s+(.+)$", raw.strip())
        if m:
            if current:
                sections[current] = buf
            current = m.group(1).strip()
            buf = []
        else:
            buf.append(raw)
    if current:
        sections[current] = buf
    return sections


def clean_lines(lines, template_hints=()):
    """清洗 markdown → 纯文本段落列表。template_hints = 模板指导行前缀元组(这些行要剔除)"""
    out = []
    for raw in lines:
        line = raw.strip()
        if not line:
            out.append("")
            continue
        if line == "---":
            continue
        if template_hints and line.startswith(template_hints):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"^>\s*", "", line)
        line = re.sub(r"^[-•]\s+", "", line)
        line = line.strip()
        out.append(line)
    while out and out[0] == "":
        out.pop(0)
    while out and out[-1] == "":
        out.pop()
    return out


def set_cell_text(cell, lines):
    """清空 cell 后写入多段文本（保留 cell 样式）。注意会删除 cell 原有全部段落。"""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run._element.getparent().remove(run._element)
    if lines:
        first.add_run(lines[0])
        for ln in lines[1:]:
            p = cell.add_paragraph()
            p.add_run(ln)


def independent_cells(row):
    """🔴 关键: row.cells 会把合并单元格重复返回多次，必须按 _tc 去重，
    否则 set_cell_text 会覆盖同一个合并区多次（最后一次覆盖前面的内容）。"""
    result = []
    for cell in row.cells:
        if not any(cell._tc is s for s in [c._tc for c in result]):
            result.append(cell)
    return result


def main():
    if len(sys.argv) != 5:
        print("用法: fill-docx-from-md.py <模板.docx> <内容.md> <输出.docx> <板块映射.json>")
        sys.exit(1)
    tmpl, md_path, out, map_path = sys.argv[1:5]
    with open(map_path, encoding="utf-8") as f:
        section_to_row = json.load(f)

    sections = read_sections(md_path)
    doc = Document(tmpl)
    table = doc.tables[0]  # 假设内容都在第一个表格（表单类模板通常是单表）

    for sec_name, row_idx in section_to_row.items():
        if sec_name not in sections:
            print(f"[SKIP] 未找到板块: {sec_name}")
            continue
        lines = clean_lines(sections[sec_name])
        row = table.rows[row_idx]
        cells = independent_cells(row)
        set_cell_text(cells[1], lines)  # cells[0]=栏目名标签, cells[1]=内容合并区
        print(f"   {sec_name}: {len(lines)} 行")

    doc.save(out)
    print(f"✅ 已生成: {out}")


if __name__ == "__main__":
    main()
